"""
QUALIFICADOR DE CEDULAS - aplicativo unico
============================================
Decodifica um .PRN, qualifica emitentes/avalistas e gera o Word --
tudo em um só programa, com uma janela simples para escolher o arquivo.

Para rodar direto:
    python qualificador_app.py

Para transformar em .exe (rodar no Windows, com Python instalado):
    pip install python-docx pyinstaller
    pyinstaller --onefile --windowed --name Qualificador qualificador_app.py

O executavel fica em dist\\Qualificador.exe -- pode copiar esse arquivo
sozinho para qualquer pasta/computador (nao precisa mais do Python nem
do Node instalado para RODAR, só para GERAR o .exe da primeira vez).
"""
import os
import re
import sys
import traceback
import tkinter as tk
from tkinter import filedialog, messagebox, scrolledtext, ttk

try:
    import openpyxl
except ImportError:
    openpyxl = None

from docx import Document
from docx.shared import Cm, Pt, Emu, RGBColor
from docx.enum.text import WD_LINE_SPACING, WD_TAB_ALIGNMENT, WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ======================================================================
# ETAPA 1: DECODIFICACAO DO PRN
# ======================================================================

MAPA_OVERSTRIKE = {
    (b'a', b'~'): 'ã', (b'A', b'~'): 'Ã',
    (b'o', b'~'): 'õ', (b'O', b'~'): 'Õ',
    (b'c', b','): 'ç', (b'C', b','): 'Ç',
    (b'a', b"'"): 'á', (b'A', b"'"): 'Á',
    (b'e', b"'"): 'é', (b'E', b"'"): 'É',
    (b'i', b"'"): 'í', (b'I', b"'"): 'Í',
    (b'o', b"'"): 'ó', (b'O', b"'"): 'Ó',
    (b'u', b"'"): 'ú', (b'U', b"'"): 'Ú',
    (b'e', b'^'): 'ê', (b'E', b'^'): 'Ê',
    (b'o', b'^'): 'ô', (b'O', b'^'): 'Ô',
    (b'a', b'^'): 'â',
    (b'a', b'`'): 'à',
}


def decodificar_prn(caminho, codepage="cp850"):
    with open(caminho, "rb") as f:
        raw = f.read()

    sem_controle = re.sub(rb'\x1bC.', b'', raw)          # ESC C <n> (3 bytes, com parametro)
    sem_controle = re.sub(rb'\x1b[A-Za-z]', b'', sem_controle)  # demais ESC + 1 letra
    sem_controle = re.sub(rb'[\x0f\x12]', b'', sem_controle)    # SI / DC2

    texto = sem_controle.decode(codepage, errors="replace")

    def resolver(m):
        base, marca = m.group(1).encode(), m.group(2).encode()
        return MAPA_OVERSTRIKE.get((base, marca), m.group(1))

    texto = re.sub(r'(.)\x08(.)', resolver, texto)
    return texto


# ======================================================================
# ETAPA 2: REESTRUTURACAO (Interveniente Garantidor -> dentro de "Por aval...")
# ======================================================================

def reestruturar_documento(texto):
    padrao_span = re.compile(
        r"INTERVENIENTE\(S\)\s+GARANTIDOR\(ES\):.*?(?=CARTILHA\s+DO\s+CR[EÉ]DITO\s+RURAL:)",
        re.IGNORECASE | re.DOTALL,
    )
    m = padrao_span.search(texto)
    if not m:
        return texto

    span_original = m.group(0)
    m_intro = re.search(r"(Assino\(amos\).*?pelo\s+emitente\.)", span_original, re.IGNORECASE | re.DOTALL)
    paragrafo_intro = m_intro.group(1) if m_intro else ""

    padrao_continua = re.compile(r"Continua\s+Proxima\s+Pagina", re.IGNORECASE)
    padrao_continuacao_header = re.compile(
        r"Continua[çc][aã]o\s+do\s+instrumento\s+de\s+cr[eé]dito\s+do\s+t[ií]tulo", re.IGNORECASE
    )

    linhas_mantidas = []
    for linha in span_original.split("\n"):
        if padrao_continua.search(linha) or padrao_continuacao_header.search(linha) or linha.strip() == "":
            linhas_mantidas.append(linha)
    span_limpo = "\n".join(linhas_mantidas)

    texto = texto.replace(span_original, span_limpo, 1)
    texto = re.sub(
        r"Por\s+aval\s+ao\(s\)\s+emitente\(s\):",
        "POR AVAL AO(S) EMITENTE(S) / INTERVENIENTE(S) GARANTIDOR(ES):\r\n      \r\n      " + paragrafo_intro,
        texto, count=1, flags=re.IGNORECASE,
    )
    return texto


# ======================================================================
# ETAPA 3: BLOCOS DE DADOS COMPLETOS (cabecalho do documento)
# ======================================================================

ROTULOS_BLOCO_DADOS = [
    ("EMITENTE", r"EMITENTE\(S\):"),
    ("AVALISTA", r"Avalista\(s\):"),
    ("CONJUGE_AVALISTA", r"C[oô]njuge\s+do\s+Avalista:"),
    ("FIDUCIANTE", r"Propriet[aá]rio\(s\):"),
]

CPF_RE = re.compile(r"CPF\.{0,20}:?\s*(?:sob\s*n[°.]?\s*)?([\d.\-]{11,})", re.IGNORECASE)
CNPJ_RE = re.compile(r"CNPJ\.{0,20}:?\s*([\d./\-]{14,})", re.IGNORECASE)
RG_RE = re.compile(r"\bRG\s*([\w./-]+)\s*-\s*([A-Z/]+)", re.IGNORECASE)
NACIONALIDADE_RE = re.compile(r"Nacionalidade\s+([A-ZÀ-Ú]+)", re.IGNORECASE)
ESTADO_CIVIL_RE = re.compile(r"Nacionalidade\s+\S+,\s*([A-ZÀ-Ú]+)", re.IGNORECASE)
FILIACAO_RE = re.compile(r"filho\(a\)\s*de\s*(.+?)\s+e\s+(.+?),", re.IGNORECASE | re.DOTALL)
PROFISSAO_RE = re.compile(
    r"(?:BENS|UNIVERSAL(?:\s+DE\s+BENS)?),\s*filho\(a\).+?,\s*(.+?),\s*(?:residente|Nacionalidade)",
    re.IGNORECASE | re.DOTALL,
)
ENDERECO_RE = re.compile(r"residente\s+e\s+domiciliado\(a\)\s+no\(a\)\s*(.+?),\s*bairro\s*(.+?),", re.IGNORECASE | re.DOTALL)
# PJ usa "com sede no(a)..." em vez de "residente e domiciliado(a) no(a)..."
ENDERECO_PJ_RE = re.compile(r"com\s+sede\s+no\(a\)\s*(.+?),\s*bairro\s*(.+?),", re.IGNORECASE | re.DOTALL)
TELEFONE_RE = re.compile(r"telefone\s*(.+?),", re.IGNORECASE | re.DOTALL)
EMAIL_RE = re.compile(r"endere[çc]o\s+eletr[ôo]nico\s*(.+?)\.\s*$", re.IGNORECASE | re.DOTALL)


def _limpar_espacos(valor):
    return re.sub(r"\s+", " ", valor).strip() if valor else None


def _fim_do_paragrafo(texto, inicio, limite):
    m = re.search(r"\r?\n[ \t]*\r?\n", texto[inicio:limite])
    return inicio + m.start() if m else limite


def _extrair_campos(bloco_completo):
    cpf_m = CPF_RE.search(bloco_completo)
    cnpj_m = CNPJ_RE.search(bloco_completo)
    eh_pj = cnpj_m is not None and cpf_m is None
    rg_m = RG_RE.search(bloco_completo)
    nac_m = NACIONALIDADE_RE.search(bloco_completo)
    ec_m = ESTADO_CIVIL_RE.search(bloco_completo)
    fil_m = FILIACAO_RE.search(bloco_completo)
    prof_m = PROFISSAO_RE.search(bloco_completo)
    end_m = ENDERECO_RE.search(bloco_completo) or (ENDERECO_PJ_RE.search(bloco_completo) if eh_pj else None)
    tel_m = TELEFONE_RE.search(bloco_completo)
    email_m = EMAIL_RE.search(bloco_completo)
    return {
        "eh_pj": eh_pj,
        "cpf": cpf_m.group(1) if cpf_m else None,
        "cnpj": cnpj_m.group(1) if cnpj_m else None,
        "rg": rg_m.group(1) if rg_m else None,
        "orgao_emissor": rg_m.group(2) if rg_m else None,
        "nacionalidade": nac_m.group(1) if nac_m else None,
        "estado_civil": ec_m.group(1) if ec_m else None,
        "pai": _limpar_espacos(fil_m.group(1)) if fil_m else None,
        "mae": _limpar_espacos(fil_m.group(2)) if fil_m else None,
        "profissao": _limpar_espacos(prof_m.group(1)) if prof_m else None,
        "endereco": _limpar_espacos(end_m.group(1)) if end_m else None,
        "bairro": _limpar_espacos(end_m.group(2)) if end_m else None,
        "telefone": _limpar_espacos(tel_m.group(1)) if tel_m else None,
        "email": _limpar_espacos(email_m.group(1)) if email_m else None,
    }


def extrair_blocos_dados(texto):
    marcadores = []
    for papel, padrao in ROTULOS_BLOCO_DADOS:
        for m in re.finditer(padrao + r"(?=[ \t]+\S)", texto):
            marcadores.append((m.start(), papel, m.group(0)))
    marcadores.sort()

    blocos = []
    for idx, (pos, papel, rotulo_txt) in enumerate(marcadores):
        limite = marcadores[idx + 1][0] if idx + 1 < len(marcadores) else pos + 1000
        fim = _fim_do_paragrafo(texto, pos, limite)
        bloco_completo = texto[pos:fim].strip()

        campos = _extrair_campos(bloco_completo)
        nome_m = re.search(re.escape(rotulo_txt.rstrip()) + r"\s*(.+?),", bloco_completo, re.IGNORECASE)
        if not campos["cpf"] and not nome_m:
            continue

        bloco_sem_rotulo = re.sub(r"^" + re.escape(rotulo_txt.rstrip()) + r"\s*", "", bloco_completo)

        blocos.append({
            "papel": papel,
            "nome": nome_m.group(1).strip() if nome_m else None,
            **campos,
            "bloco_completo": bloco_completo,
            "bloco_sem_rotulo": bloco_sem_rotulo,
        })
    return blocos


# ======================================================================
# ETAPA 4: CAMPO DE ASSINATURAS
# ======================================================================

CABECALHOS_ASSINATURA = [
    ("EMITENTE", r"^[ \t]*EMITENTE\(S\):[ \t]*\r?$"),
    ("AVALISTA", r"^[ \t]*Por\s+aval\s+ao\(s\)\s+emitente\(s\)(?:\s*/\s*INTERVENIENTE\(S\)\s+GARANTIDOR\(ES\))?:"),
    ("INTERVENIENTE_GARANTIDOR", r"^[ \t]*INTERVENIENTE\(S\)\s+GARANTIDOR\(ES\):"),
]

CPF_ASSINATURA_RE = re.compile(r"\bCPF\.{0,20}:?\s*([\d.\-]{11,})", re.IGNORECASE)
AUTORIZACAO_CONJUGE_RE = re.compile(
    r"Autoriza[çc][aã]o\s+para\s+os\s+fins\s+do\s*\r?\n?\s*Art\.\s*1\.647\s+do\s+C[oó]digo\s+Civil",
    re.IGNORECASE,
)
# rotulo "Interveniente(s) Garantidor(es)" SOZINHO (sem dois pontos no
# fim), que pode aparecer no meio de outro bloco (ex: dentro do
# EMITENTE(S), quando quem assina como interveniente nao e o emitente).
# Diferente de CABECALHOS_ASSINATURA (que exige ":" no fim e vira o
# papel "oficial" da secao), esse so troca o papel DAQUELA assinatura
# especifica, sem mudar o rastreamento do resto do documento.
PADRAO_INTERVENIENTE_LABEL_SOLTO = re.compile(
    r"^[ \t]*Interveniente\(s\)\s+Garantidor\(es\)[ \t]*:?[ \t]*$", re.IGNORECASE | re.MULTILINE
)


PADRAO_RAZAO_SOCIAL = re.compile(r"Raz[ãa]o\s+Social:\s*(.+)", re.IGNORECASE)


def _proxima_linha_assinatura(linhas_flat, indice):
    """A partir de indice, pula linhas em branco e marcadores de
    paginacao (a assinatura pode atravessar quebra de pagina -- ex:
    'Razao Social:' numa pagina e o 'CNPJ' na seguinte) procurando a
    proxima linha com conteudo real. Para e devolve (None, indice) se
    encontrar OUTRO sublinhado (comeco de nova assinatura)."""
    j = indice
    while j < len(linhas_flat):
        texto_j = linhas_flat[j]["texto"]
        if PADRAO_UNDERLINE.match(texto_j):
            return None, j
        eh_pag = (
            PADRAO_CONTINUA_PROXIMA.search(texto_j)
            or PADRAO_CONTINUACAO_HEADER.search(texto_j)
            or PADRAO_PAGINA_SOZINHA.match(texto_j)
        )
        if texto_j.strip() == "" or eh_pag:
            j += 1
            continue
        return linhas_flat[j], j
    return None, j


def extrair_assinaturas(texto):
    marcadores_papel = []
    for papel, padrao in CABECALHOS_ASSINATURA:
        for m in re.finditer(padrao, texto, re.IGNORECASE | re.MULTILINE):
            marcadores_papel.append((m.start(), papel))
    marcadores_papel.sort()

    def papel_na_posicao(pos):
        papel = None
        for p, nome_papel in marcadores_papel:
            if p <= pos:
                papel = nome_papel
            else:
                break
        return papel

    # junta as duas formas de assinatura -- PF ("NOME:") e PJ ("Razao
    # Social:") -- numa unica lista ordenada por posicao no texto
    ocorrencias = []
    for m in re.finditer(r"\bNOME:\s*(.+)", texto, re.IGNORECASE):
        ocorrencias.append((m.start(), False, m.group(1).strip()))
    for m in re.finditer(PADRAO_RAZAO_SOCIAL, texto):
        ocorrencias.append((m.start(), True, m.group(1).strip()))
    ocorrencias.sort(key=lambda o: o[0])

    assinaturas = []
    for pos, eh_pj, nome in ocorrencias:
        janela = texto[pos:pos + 250]
        if eh_pj:
            doc_m = CNPJ_RE.search(janela)
        else:
            doc_m = CPF_ASSINATURA_RE.search(janela)
        documento = doc_m.group(1) if doc_m else None
        contexto_anterior = texto[max(0, pos - 150):pos]
        eh_autorizacao_conjuge = bool(AUTORIZACAO_CONJUGE_RE.search(contexto_anterior))
        eh_interveniente_local = bool(PADRAO_INTERVENIENTE_LABEL_SOLTO.search(contexto_anterior))

        papel = papel_na_posicao(pos)
        if eh_interveniente_local and not eh_autorizacao_conjuge:
            papel = "INTERVENIENTE_GARANTIDOR"

        assinaturas.append({
            "papel_assinatura": papel,
            "autorizacao_conjuge_art_1647": eh_autorizacao_conjuge,
            "eh_pj": eh_pj,
            "nome": nome,
            "cpf": None if eh_pj else documento,
            "cnpj": documento if eh_pj else None,
        })
    return assinaturas


def papel_header_esperado(assinatura):
    if assinatura["autorizacao_conjuge_art_1647"]:
        return "CONJUGE_AVALISTA"
    return assinatura["papel_assinatura"]  # INTERVENIENTE_GARANTIDOR cai no fallback


def qualificar_com_blocos(assinaturas, blocos):
    blocos_por_doc_e_papel = {}
    blocos_por_doc_qualquer = {}
    for b in blocos:
        documento = b.get("cnpj") or b.get("cpf")
        if not documento:
            continue
        blocos_por_doc_e_papel[(documento, b["papel"])] = b
        blocos_por_doc_qualquer.setdefault(documento, b)

    qualificadas = []
    for assinatura in assinaturas:
        documento = assinatura.get("cnpj") or assinatura.get("cpf")
        papel_esperado = papel_header_esperado(assinatura)
        dados_exatos = blocos_por_doc_e_papel.get((documento, papel_esperado))
        dados_completos = dados_exatos or blocos_por_doc_qualquer.get(documento)
        qualificadas.append({
            **assinatura,
            "encontrado_no_cabecalho": dados_completos is not None,
            "match_exato_de_papel": dados_exatos is not None,
            "dados_completos": dados_completos,
        })
    return qualificadas


# ======================================================================
# ETAPA 4.5: SECAO "2 - IMOVEIS" -- deteccao de hipoteca e avaliacao
# ======================================================================

PADRAO_SECAO_IMOVEIS = re.compile(r"2\s*-\s*IM[OÓ]VEIS:", re.IGNORECASE)
PADRAO_FIM_SECAO_IMOVEIS = re.compile(
    r"CONSTITUI[ÇC][ÃA]O\s+DA\s+ALIENA[ÇC][ÃA]O\s+FIDUCI[ÁA]RIA|II\s*-\s*CL[ÁA]USULAS",
    re.IGNORECASE,
)
PADRAO_HIPOTECA = re.compile(r"\bHIPOTEC\w*\b", re.IGNORECASE)
PADRAO_AVALIADO_EM = re.compile(r"Avaliado\s+em\s+([\d.,]+)\s*\(\s*(.*?)\s*\)", re.IGNORECASE | re.DOTALL)
PADRAO_VALOR_GARANTIA = re.compile(
    r"Valor\s+de\s+avalia[çc][ãa]o\s+do\s+im[óo]vel\s+para\s+fins\s+de\s+garantia\s+e\s+venda\s+em\s+p[úu]blico\s+leil[ãa]o:\s*R\$\s*([\d.,]+)\s*\(\s*(.*?)\s*\)",
    re.IGNORECASE | re.DOTALL,
)


def _normalizar_espacos(s):
    return re.sub(r"\s+", " ", s).strip()


PADRAO_TEXTO_COPIA_MATRICULA = re.compile(
    r"A\s+c[óo]pia\s+das\s+referidas\s+matr[íi]culas\s+fazem\s+parte\s+integrante\s+e\s+"
    r"insepar[áa]vel\s+desta\s+C[ée]dula,\s*para\s+todos\s+os\s+fins\s+e\s+efeitos,\s*"
    r"como\s+se\s+aqui\s+estivesse\s+integralmente\s+transcritas\.",
    re.IGNORECASE | re.DOTALL,
)


def _remover_boilerplate_paginacao(bloco):
    """Remove as linhas de paginacao (Continua Proxima Pagina /
    Continuacao do instrumento... / Pagina: N sozinha) de dentro de um
    trecho de texto -- usado pra descricao do imovel nao "puxar" esses
    artefatos quando a secao atravessa mais de uma pagina do PRN. Tambem
    remove a frase fixa "A copia das referidas matriculas..." que nao e
    parte da descricao do imovel em si."""
    bloco = PADRAO_TEXTO_COPIA_MATRICULA.sub("", bloco)
    linhas = bloco.split("\n")
    linhas_limpas = [
        l for l in linhas
        if not (
            PADRAO_CONTINUA_PROXIMA.search(l)
            or PADRAO_CONTINUACAO_HEADER.search(l)
            or PADRAO_PAGINA_SOZINHA.match(l)
        )
    ]
    return "\n".join(linhas_limpas)


def detectar_secao_imoveis(texto):
    """
    Localiza a secao "2 - IMOVEIS:" e verifica: (a) se ha mencao a
    hipoteca dentro dela, (b) os valores de avaliacao informados, e
    (c) se os dois valores de avaliacao sao identicos (possivel erro
    de preenchimento duplicado). Retorna None se a secao nao existir
    no documento (ex: cedulas sem esse tipo de garantia).

    A secao pode atravessar varias paginas do PRN (descricao de imovel
    grande) -- o limite de seguranca e bem folgado (50 mil caracteres)
    pra nao cortar nada nesse caso, e as linhas de paginacao que
    aparecem no meio do caminho sao removidas antes de analisar.
    """
    m_inicio = PADRAO_SECAO_IMOVEIS.search(texto)
    if not m_inicio:
        return None

    m_fim = PADRAO_FIM_SECAO_IMOVEIS.search(texto, m_inicio.end())
    fim = m_fim.start() if m_fim else min(m_inicio.end() + 50000, len(texto))
    bloco = _remover_boilerplate_paginacao(texto[m_inicio.start():fim])

    hipoteca_m = PADRAO_HIPOTECA.search(bloco)
    avaliado_m = PADRAO_AVALIADO_EM.search(bloco)
    garantia_m = PADRAO_VALOR_GARANTIA.search(bloco)

    valores = []
    if avaliado_m:
        valores.append({
            "rotulo": "Avaliado em",
            "valor": avaliado_m.group(1).strip(),
            "extenso": _normalizar_espacos(avaliado_m.group(2)),
        })
    if garantia_m:
        valores.append({
            "rotulo": "Valor de avaliação para fins de garantia e venda em público leilão",
            "valor": garantia_m.group(1).strip(),
            "extenso": _normalizar_espacos(garantia_m.group(2)),
        })

    duplicado = (
        len(valores) == 2
        and valores[0]["valor"].replace(" ", "") == valores[1]["valor"].replace(" ", "")
    )

    return {
        "span_inicio": m_inicio.start(),
        "span_fim": fim,
        "bloco_original": bloco,
        "tem_hipoteca": hipoteca_m is not None,
        "trecho_hipoteca": _normalizar_espacos(bloco[max(0, hipoteca_m.start() - 60):hipoteca_m.end() + 60]) if hipoteca_m else None,
        "valores_avaliacao": valores,
        "avaliacoes_duplicadas": duplicado,
    }


# --------------------------------------------------------------------
# Hipoteca Cedular -- mesma ideia da secao "2 - IMOVEIS", mas pro outro
# tipo de garantia (cedula rural hipotecaria). O PRN as vezes duplica um
# trecho inteiro dentro dessa clausula (bug de geracao do lado deles) --
# detectamos e removemos a repeticao automaticamente.
# --------------------------------------------------------------------

PADRAO_SECAO_HIPOTECA = re.compile(
    r"HIPOTECA\s+CEDULAR\s*-\s*Em\s+seguran[çc]a\s+das\s+obriga[çc][õo]es\s+contratadas,",
    re.IGNORECASE,
)
PADRAO_FIM_SECAO_HIPOTECA = re.compile(
    r"Al[ée]m\s+das\s+declara[çc][õo]es\s+j[áa]\s+prestadas", re.IGNORECASE
)


def _remover_duplicacao_consecutiva(texto, tamanho_minimo=30):
    """Detecta se um trecho substancial de texto se repete duas vezes
    SEGUIDAS (bug conhecido na geracao do PRN pelo sistema de origem --
    ex: endereco do avalista duplicado dentro da clausula de hipoteca) e
    remove a repeticao, mantendo so uma ocorrencia. So mexe se achar
    uma repeticao de verdade; se nao achar, devolve o texto igual."""
    palavras = texto.split()
    n = len(palavras)
    melhor = None
    for tam in range(n // 2, 3, -1):
        for inicio in range(0, n - 2 * tam + 1):
            bloco1 = palavras[inicio:inicio + tam]
            bloco2 = palavras[inicio + tam:inicio + 2 * tam]
            if bloco1 == bloco2 and len(" ".join(bloco1)) >= tamanho_minimo:
                melhor = (inicio, tam)
                break
        if melhor:
            break
    if not melhor:
        return texto
    inicio, tam = melhor
    novas_palavras = palavras[:inicio + tam] + palavras[inicio + 2 * tam:]
    return " ".join(novas_palavras)


def detectar_secao_hipoteca(texto):
    """Localiza a clausula 'HIPOTECA CEDULAR - Em seguranca das
    obrigacoes contratadas, ...' e devolve o trecho (ja sem duplicacao,
    se detectada) pronto pra revisao/edicao futura. Retorna None se a
    cedula nao tiver esse tipo de garantia."""
    m_inicio = PADRAO_SECAO_HIPOTECA.search(texto)
    if not m_inicio:
        return None

    m_fim = PADRAO_FIM_SECAO_HIPOTECA.search(texto, m_inicio.end())
    fim = m_fim.start() if m_fim else min(m_inicio.end() + 50000, len(texto))
    bloco_bruto = _remover_boilerplate_paginacao(texto[m_inicio.start():fim])

    bloco_sem_duplicacao = _remover_duplicacao_consecutiva(bloco_bruto)
    tinha_duplicacao = bloco_sem_duplicacao != bloco_bruto

    return {
        "span_inicio": m_inicio.start(),
        "span_fim": fim,
        "bloco_original": bloco_bruto,
        "bloco_sem_duplicacao": bloco_sem_duplicacao,
        "tinha_duplicacao": tinha_duplicacao,
    }


def aplicar_correcao_hipoteca(texto, deteccao):
    """Substitui o trecho original da clausula de hipoteca pela versao
    sem duplicacao (mantem o mesmo recuo/formatacao do resto do
    documento)."""
    linhas_novas = []
    for linha in deteccao["bloco_sem_duplicacao"].split("\n"):
        linhas_novas.append("      " + linha.strip() if linha.strip() else "      ")
    bloco_novo = "\r\n".join(linhas_novas)
    return texto[:deteccao["span_inicio"]] + bloco_novo + texto[deteccao["span_fim"]:]


# --------------------------------------------------------------------
# Aba "Assinaturas" -- visualizacao/edicao de tudo que sai como
# qualificacao de cada assinatura (emitente/avalista/PJ), num unico
# texto editavel, delimitado por marcadores que dao pra editar,
# remover (deixar em branco) ou acrescentar uma assinatura nova.
# --------------------------------------------------------------------

PADRAO_MARCADOR_ASSINATURA_PREVIEW = re.compile(
    r"^-{3}\s*(\d+)\.\s*([^(]+?)\s*\(([^)]*)\)\s*-{3}\s*$", re.MULTILINE
)


def montar_texto_revisao_assinaturas(assinaturas_qualificadas):
    """Monta o texto editavel da aba 'Assinaturas': um bloco por
    assinatura, com o papel/documento no marcador e o texto completo
    (o mesmo que vai ser usado na qualificacao) editavel embaixo."""
    partes = []
    for i, a in enumerate(assinaturas_qualificadas, start=1):
        documento = a.get("cnpj") or a.get("cpf") or "sem documento"
        papel = a["papel_assinatura"]
        if a.get("autorizacao_conjuge_art_1647"):
            papel += " - AUTORIZAÇÃO CÔNJUGE"
        if a.get("eh_pj"):
            papel += " (PJ)"
        dados = a.get("dados_completos")
        if dados and dados.get("bloco_sem_rotulo"):
            linhas_originais = dados["bloco_sem_rotulo"].split("\r\n") if "\r\n" in dados["bloco_sem_rotulo"] else dados["bloco_sem_rotulo"].split("\n")
            texto_bloco = _reflow_prosa(linhas_originais)
        else:
            texto_bloco = f"(NÃO ENCONTRADO NO CABEÇALHO -- nome: {a.get('nome') or '?'})"
        partes.append(f"--- {i}. {papel} ({documento}) ---\n{texto_bloco}")
    partes.append(
        "--- NOVO: NOME DO PAPEL (documento) ---\n"
        "(pra acrescentar uma assinatura que não veio no PRN, copie um bloco acima,\n"
        "troque 'NOVO' por um número maior que os existentes, e preencha os dados)"
    )
    return "\n\n".join(partes)


def analisar_texto_revisao_assinaturas(texto_editado):
    """Le o texto (possivelmente editado) da aba 'Assinaturas' e devolve
    uma lista ordenada de {indice, papel, documento, texto}. Blocos com
    indice 'NOVO' ou nao numerico sao ignorados (o texto de exemplo no
    fim), a nao ser que o colaborador tenha trocado por um numero."""
    marcadores = list(PADRAO_MARCADOR_ASSINATURA_PREVIEW.finditer(texto_editado))
    resultado = []
    for idx, m in enumerate(marcadores):
        inicio_texto = m.end()
        fim_texto = marcadores[idx + 1].start() if idx + 1 < len(marcadores) else len(texto_editado)
        corpo = texto_editado[inicio_texto:fim_texto].strip()
        try:
            indice_original = int(m.group(1))
        except ValueError:
            continue
        resultado.append({
            "indice_original": indice_original,
            "papel": m.group(2).strip(),
            "documento": m.group(3).strip(),
            "texto": corpo,
        })
    return resultado


def aplicar_revisao_assinaturas(texto, assinaturas_qualificadas, texto_editado, log=print):
    """Aplica as edicoes da aba 'Assinaturas' -- sobrescreve o texto de
    qualificacao de cada assinatura existente com o que foi editado, e
    acrescenta como novas assinaturas (apos a ultima existente) qualquer
    bloco com indice alem da quantidade original. Devolve o texto
    (com as assinaturas novas inseridas, se houver -- as edicoes das
    existentes sao aplicadas depois, sobre assinaturas_qualificadas)."""
    entradas = analisar_texto_revisao_assinaturas(texto_editado)
    total_original = len(assinaturas_qualificadas)

    for entrada in entradas:
        idx = entrada["indice_original"] - 1
        if 0 <= idx < total_original:
            dc = assinaturas_qualificadas[idx].get("dados_completos")
            if dc is None:
                dc = {}
                assinaturas_qualificadas[idx]["dados_completos"] = dc
            dc["bloco_sem_rotulo"] = entrada["texto"]

    novas = [e for e in entradas if e["indice_original"] > total_original]
    for nova in novas:
        log(f"  Acrescentando assinatura nova: {nova['papel']} ({nova['documento']}) ...")
        bloco = (
            "\r\n      \r\n"
            + "      " + "_" * 50 + "\r\n"
            + "      " + MARCADOR_ADICIONADO_INICIO + nova["texto"] + MARCADOR_ADICIONADO_FIM + "\r\n"
        )
        pos = _posicao_apos_ultima_assinatura(texto)
        if pos is not None:
            texto = texto[:pos] + bloco + texto[pos:]

    return texto


def montar_texto_revisao(deteccao):
    """Monta o texto EDITAVEL mostrado na tela de revisao: tudo que vem
    depois de "2 - IMOVEIS:" ate onde a secao termina (onde a avaliacao
    e encontrada) -- reflui as quebras de linha de largura fixa do PRN
    em paragrafos corridos, mais faceis de ler/editar na caixa de texto."""
    conteudo = PADRAO_SECAO_IMOVEIS.sub("", deteccao["bloco_original"], count=1)

    paragrafos_brutos = re.split(r"\r?\n[ \t]*\r?\n", conteudo)
    paragrafos = []
    for p in paragrafos_brutos:
        linha_unica = _normalizar_espacos(p.replace("\r\n", " ").replace("\n", " "))
        if linha_unica:
            paragrafos.append(linha_unica)

    aviso_linhas = []
    if deteccao["tem_hipoteca"]:
        aviso_linhas.append(f"[AVISO: HIPOTECA ENCONTRADA -- trecho: {deteccao['trecho_hipoteca']}]")
        aviso_linhas.append("")

    return "\n".join(aviso_linhas) + "\n\n".join(paragrafos)


def aplicar_correcao_imoveis(texto, deteccao, texto_revisado):
    """Substitui o bloco original da secao '2 - IMOVEIS:' pelo texto
    revisado (editado na tela), formatado com o mesmo recuo padrao do
    resto do documento. Devolve (novo_texto, posicao_onde_a_descricao_termina)
    -- essa posicao e usada por aplicar_clausula_superveniencia para
    inserir a clausula logo apos a descricao, mesmo que ela tenha mudado
    de tamanho."""
    linhas_novas = ["      2 - IMÓVEIS:", "      "]
    for linha in texto_revisado.split("\n"):
        linhas_novas.append("      " + linha if linha.strip() else "      ")
    bloco_novo = "\r\n".join(linhas_novas) + "\r\n      \r\n"

    novo_texto = texto[:deteccao["span_inicio"]] + bloco_novo + texto[deteccao["span_fim"]:]
    posicao_fim = deteccao["span_inicio"] + len(bloco_novo)
    return novo_texto, posicao_fim


# --------------------------------------------------------------------
# Clausula de Superveniencia (aba "Opcoes")
# --------------------------------------------------------------------

TEXTO_CLAUSULA_SUPERVENIENCIA = (
    "CONSTITUIÇÃO DA ALIENAÇÃO FIDUCIÁRIA DA PROPRIEDADE SUPERVENIENTE: "
    "O imóvel objeto da garantia é constituído em ALIENAÇÃO FIDUCIÁRIA SUPERVENIENTE, "
    "nos termos do §3º do art. 22 da Lei n° 9.514/1997, com a redação dada pela Lei "
    "14.711/2023. Parágrafo Único: Na hipótese de constituição de alienações "
    "fiduciárias sucessivas da propriedade superveniente e/ou quando houver a extensão "
    "da garantia fiduciária, havendo o inadimplemento de quaisquer das obrigações "
    "garantidas pelo mesmo imóvel, fica desde já autorizado ao CREDOR determinar o "
    "vencimento antecipado das demais obrigações de que for titular, nos termos do §6º "
    "do artigo 22 da Lei 9.514/2023, com a redação dada pela Lei 14.711/2023."
)

PADRAO_CLAUSULA_SUPERVENIENCIA = re.compile(
    r"CONSTITUI[ÇC][ÃA]O\s+DA\s+ALIENA[ÇC][ÃA]O\s+FIDUCI[ÁA]RIA\s+DA\s+PROPRIEDADE\s+SUPERVENIENTE",
    re.IGNORECASE,
)

# marcadores invisiveis (area de uso privado do Unicode -- nao aparecem
# como texto de verdade) usados so pra sinalizar pro gerador do Word
# qual trecho foi ACRESCENTADO pelo script, pra colorir diferente do
# resto e o colaborador saber o que foi adicionado x o que veio no PRN.
MARCADOR_ADICIONADO_INICIO = "\ue000"
MARCADOR_ADICIONADO_FIM = "\ue001"

# modelo de representacao por socio, acrescentado apos os dados de uma
# assinatura PJ -- o PRN nao traz nome/CPF do socio, entao fica com
# asteriscos pro colaborador preencher manualmente (destacado em azul,
# igual ao resto do que o script acrescenta)
TEXTO_REPRESENTACAO_PJ = (
    "Neste ato é representada por sua sócios **********, inscrita no CPF nº "
    "**********, conforme disposto no contrato social."
)


def detectar_clausula_superveniencia(texto):
    """True se a clausula de superveniencia ja existe no documento."""
    return bool(PADRAO_CLAUSULA_SUPERVENIENCIA.search(texto))


def aplicar_clausula_superveniencia(texto, posicao_insercao, adicionar):
    """Insere a clausula de superveniencia logo na posicao indicada (logo
    apos a descricao do imovel), marcada como "acrescentada" (cor
    diferente no Word). So insere se adicionar=True -- quem decide ANTES
    se deve chamar isso ou nao e a logica de "so acrescenta se nao
    existia e o colaborador marcou o checkbox" (ver finalizar_geracao)."""
    if not adicionar:
        return texto

    bloco = (
        "\r\n      \r\n      "
        + MARCADOR_ADICIONADO_INICIO
        + TEXTO_CLAUSULA_SUPERVENIENCIA
        + MARCADOR_ADICIONADO_FIM
        + "\r\n      \r\n"
    )
    return texto[:posicao_insercao] + bloco + texto[posicao_insercao:]


PADRAO_FIM_CLAUSULA_SUPERVENIENCIA = re.compile(r"II\s*-\s*CL[ÁA]USULAS", re.IGNORECASE)


def remover_clausula_superveniencia(texto):
    """Remove a clausula de superveniencia que ja existia no documento
    original -- usada quando o checkbox vem marcado (a clausula foi
    encontrada) e o colaborador desmarca, decidindo tirar ela da cedula."""
    m_inicio = PADRAO_CLAUSULA_SUPERVENIENCIA.search(texto)
    if not m_inicio:
        return texto
    m_fim = PADRAO_FIM_CLAUSULA_SUPERVENIENCIA.search(texto, m_inicio.end())
    fim = m_fim.start() if m_fim else min(m_inicio.end() + 2000, len(texto))
    return texto[:m_inicio.start()] + texto[fim:]


# --------------------------------------------------------------------
# Cartorio (aba "Opcoes") -- deteccao simples por enquanto; no futuro vai
# cruzar com uma planilha de criterios por cartorio.
# --------------------------------------------------------------------

PADRAO_CARTORIO = re.compile(
    r"Registro\s+de\s+Im[oó]veis[^\r\n,\.]*(?:Comarca\s+de\s+[^\r\n,\.]+)?",
    re.IGNORECASE,
)


def detectar_cartorio(texto):
    """Tenta achar o nome do cartorio/comarca citado na descricao do
    imovel (ex: 'Registro de Imoveis Comarca de Mandaguacu - PR').
    Retorna None se nao encontrar nada parecido."""
    m = PADRAO_CARTORIO.search(texto)
    return _normalizar_espacos(m.group(0)) if m else None


# --------------------------------------------------------------------
# Procuradores (aba "Opcoes") -- le a planilha "Procuradores -
# Cartorios.xlsx" que o colaborador mantem atualizada numa pasta
# "documentos" do lado do programa (funciona bem com OneDrive
# sincronizado nessa mesma pasta).
# --------------------------------------------------------------------

NOME_PLANILHA_PROCURADORES = "Procuradores - Cartórios.xlsx"
NOME_PASTA_DOCUMENTOS = "documentos"

TEXTO_CREDOR_BASE = (
    "CREDOR: COOPERATIVA DE CREDITO POUPANCA E INVESTIMENTO DEXIS SICREDI DEXIS, "
    "instituição financeira brasileira, CNPJ 79.342.069/0001-53 doravante denominada "
    "CREDORA, estabelecida no(a) AVENIDA PARANA 891, na cidade de MARINGA/PR. "
    "NESTE ATO REPRESENTADA POR SEUS PROCURADORES {procurador_1} E {procurador_2}"
)


def pasta_base_app():
    """Pasta onde o programa esta rodando de verdade -- funciona tanto
    rodando o .py direto quanto rodando o .exe empacotado (PyInstaller
    onefile extrai pra uma pasta temporaria, entao precisa usar
    sys.executable nesse caso, nao __file__)."""
    if getattr(sys, "frozen", False):
        return os.path.dirname(sys.executable)
    return os.path.dirname(os.path.abspath(__file__))


def pasta_documentos():
    """Pasta 'documentos' do lado do programa, onde fica a planilha de
    procuradores/cartorios. Cria a pasta se ainda nao existir, pra
    aparecer pronta pro colaborador arrastar o arquivo pra dentro."""
    caminho = os.path.join(pasta_base_app(), NOME_PASTA_DOCUMENTOS)
    try:
        os.makedirs(caminho, exist_ok=True)
    except OSError:
        pass
    return caminho


def caminho_planilha_procuradores():
    return os.path.join(pasta_documentos(), NOME_PLANILHA_PROCURADORES)


def ler_planilha_procuradores(log=print):
    """Le a planilha de procuradores/cartorios e devolve uma lista de
    agencias -- so as que tem PELO MENOS um procurador (1 ou 2)
    completamente preenchido (nome, CPF e cargo). Devolve lista vazia
    se a planilha nao existir ou o openpyxl nao estiver disponivel."""
    if openpyxl is None:
        log("  (openpyxl não instalado -- dropdown de Procuradores ficará vazio.)")
        return []

    caminho = caminho_planilha_procuradores()
    if not os.path.isfile(caminho):
        log(f"  Planilha de procuradores não encontrada em: {caminho}")
        return []

    try:
        wb = openpyxl.load_workbook(caminho, data_only=True)
        ws = wb.active
        linhas = list(ws.iter_rows(values_only=True))
    except Exception as e:
        log(f"  AVISO: não consegui ler a planilha de procuradores ({e}).")
        return []

    if not linhas:
        return []

    cabecalho = [str(c).strip() if c else "" for c in linhas[0]]

    def idx(nome_coluna):
        for i, c in enumerate(cabecalho):
            if c.lower() == nome_coluna.lower():
                return i
        return None

    col_codigo = idx("Código Agência")
    col_agencia = idx("Agência")
    col_nome1, col_cpf1, col_cargo1 = idx("Nome do Procurador - 1"), idx("CPF - 1"), idx("CARGO - 1")
    col_nome2, col_cpf2, col_cargo2 = idx("Nome do Procurador - 2"), idx("CPF - 2"), idx("CARGO - 2")

    def valor(linha, col):
        if col is None or col >= len(linha) or linha[col] is None:
            return ""
        return str(linha[col]).strip()

    agencias = []
    for linha in linhas[1:]:
        if linha is None or all(c is None for c in linha):
            continue
        codigo = valor(linha, col_codigo)
        agencia = valor(linha, col_agencia)
        proc1 = {"nome": valor(linha, col_nome1), "cpf": valor(linha, col_cpf1), "cargo": valor(linha, col_cargo1)}
        proc2 = {"nome": valor(linha, col_nome2), "cpf": valor(linha, col_cpf2), "cargo": valor(linha, col_cargo2)}
        proc1_completo = all(proc1.values())
        proc2_completo = all(proc2.values())
        if not (proc1_completo or proc2_completo):
            continue  # nenhum procurador preenchido -- nao aparece pra selecionar
        agencias.append({
            "codigo": codigo,
            "agencia": agencia,
            "rotulo": f"{codigo} - {agencia}",
            "procurador_1": proc1 if proc1_completo else None,
            "procurador_2": proc2 if proc2_completo else None,
        })
    return agencias


def montar_texto_credor_procuradores(agencia):
    """Monta a linha 'CREDOR: ...' com os dados do(s) procurador(es) da
    agencia selecionada, pronta pra entrar como assinatura apos o
    ultimo campo de assinatura da cedula."""
    def formatar(proc):
        if not proc:
            return None
        return f"{proc['nome']}, inscrito(a) no CPF nº {proc['cpf']}, {proc['cargo']}"

    p1 = formatar(agencia.get("procurador_1"))
    p2 = formatar(agencia.get("procurador_2"))
    if p1 and p2:
        texto = TEXTO_CREDOR_BASE.format(procurador_1=p1, procurador_2=p2)
    elif p1:
        texto = TEXTO_CREDOR_BASE.replace(" E {procurador_2}", "").format(procurador_1=p1)
    elif p2:
        texto = TEXTO_CREDOR_BASE.replace("{procurador_1} E ", "").format(procurador_2=p2)
    else:
        texto = TEXTO_CREDOR_BASE.replace(" NESTE ATO REPRESENTADA POR SEUS PROCURADORES {procurador_1} E {procurador_2}", "")
    return texto


def _posicao_apos_ultima_assinatura(texto):
    """Acha a posicao logo apos o NOME/Razao Social + CPF/CNPJ do
    ULTIMO campo de assinatura existente no texto -- usada tanto pra
    acrescentar a assinatura do procurador quanto novas assinaturas
    vindas da aba de revisao. Devolve None se nao achar nenhuma."""
    ultima_pos = 0
    for m in re.finditer(r"\bNOME:\s*.+|\bRaz[ãa]o\s+Social:\s*.+", texto, re.IGNORECASE):
        ultima_pos = m.end()
    if ultima_pos == 0:
        return None

    # avanca ate o fim da linha do CPF/CNPJ que segue o NOME/Razao Social
    fim_linha_doc_m = re.search(r"\r?\n", texto[ultima_pos:])
    pos_insercao = ultima_pos + fim_linha_doc_m.end() if fim_linha_doc_m else ultima_pos
    m_doc = re.search(r"\bCPF[\.\s:]*[\d.\-]+|\bCNPJ[\.\s:]*[\d./\-]+", texto[pos_insercao:pos_insercao + 200], re.IGNORECASE)
    if m_doc:
        fim_doc_linha = re.search(r"\r?\n", texto[pos_insercao + m_doc.end():])
        if fim_doc_linha:
            pos_insercao = pos_insercao + m_doc.end() + fim_doc_linha.end()
    return pos_insercao


def aplicar_assinatura_procurador(texto, agencia):
    """Acrescenta o campo de assinatura da CREDORA (representada pelos
    procuradores da agencia escolhida) logo apos o ultimo campo de
    assinatura existente na cedula, marcado em azul (acrescentado pelo
    script)."""
    pos_insercao = _posicao_apos_ultima_assinatura(texto)
    if pos_insercao is None:
        return texto

    texto_credor = montar_texto_credor_procuradores(agencia)
    bloco = (
        "\r\n      \r\n"
        + "      " + "_" * 50 + "\r\n"
        + "      " + MARCADOR_ADICIONADO_INICIO + texto_credor + MARCADOR_ADICIONADO_FIM + "\r\n"
    )
    return texto[:pos_insercao] + bloco + texto[pos_insercao:]



# ETAPA 5: GERAR O WORD (python-docx)
# ======================================================================

FATOR_LARGURA_CHAR = 0.6  # Courier New: ~0.6em de largura por caractere
FATOR_ALTURA_LINHA = 1.05
FONTE_PT = 12
FONTE_MENOR_PT = 8
MARCADOR_FONTE_MENOR = "Extrato de Operações SICOR"
RECUO_PADRAO = "      "
PADRAO_UNDERLINE = re.compile(r"^\s*_{5,}\s*$")
PADRAO_AUTORIZACAO_UMA_LINHA = re.compile(r"Autoriza[çc][aã]o\s+para\s+os\s+fins\s+do", re.IGNORECASE)
PADRAO_AUTORIZACAO_L1 = re.compile(r"Autoriza[çc][aã]o\s+para\s+os\s+fins\s+do\s*$", re.IGNORECASE)
PADRAO_AUTORIZACAO_L2 = re.compile(r"^\s*Art\.\s*1\.647\s+do\s+C[oó]digo\s+Civil", re.IGNORECASE)
PADRAO_CONTINUA_PROXIMA = re.compile(r"Continua\s+Proxima\s+Pagina", re.IGNORECASE)
PADRAO_CONTINUACAO_HEADER = re.compile(
    r"Continua[çc][aã]o\s+do\s+instrumento\s+de\s+cr[eé]dito\s+do\s+t[ií]tulo", re.IGNORECASE
)
PADRAO_PAGINA_SOZINHA = re.compile(r"^\s*Pagina:\s*\d+\s*$", re.IGNORECASE)


def _pagina_so_tem_boilerplate(linhas):
    """True se a pagina nao tem nenhum conteudo real -- só linhas em
    branco e/ou os avisos de paginacao (Continua Proxima Pagina /
    Continuacao do instrumento...)."""
    for linha in linhas:
        if linha.strip() == "":
            continue
        if PADRAO_CONTINUA_PROXIMA.search(linha) or PADRAO_CONTINUACAO_HEADER.search(linha):
            continue
        return False
    return True


def _linhas_de_qualificacao(dados):
    if not dados or not dados.get("bloco_sem_rotulo"):
        return []
    bloco = dados["bloco_sem_rotulo"]
    linhas = bloco.split("\r\n") if "\r\n" in bloco else bloco.split("\n")
    return [RECUO_PADRAO + l if i == 0 else l for i, l in enumerate(linhas)]


PADRAO_TABULAR_ESPACO = re.compile(r"\S {4,}\S")
PADRAO_SEPARADOR = re.compile(r"^[\-=_]{5,}$")
PADRAO_CAMPO_FORMULARIO = re.compile(r"\.{3,}\s*:")
PADRAO_CABECALHO_TABELA = re.compile(r"^(Nro\s+Data|Ref\.BACEN|Quadro\s+Resumo\s+da)", re.IGNORECASE)
# numeracao de clausula/item no comeco da linha (1., 3., (iii), a), i.) --
# sozinha ja cria um "gap" de varios espacos que nao e uma tabela de verdade
PADRAO_NUMERACAO_CLAUSULA = re.compile(r"^\s*(?:\(?[ivxlcdm]{1,6}\)|\(?[a-zA-Z]\)|\d+[.\)])\s+", re.IGNORECASE)

# linha de parcela (Nro, Data, Valor, Percentual) -- usado pra separar
# quando o PRN concatena mais de uma parcela na mesma linha fisica
PADRAO_LINHA_PARCELA = re.compile(
    r"(\d+)\s+(\d{2}/\d{2}/\d{4})\s+([\d.]+,\d{2})\s+([\d]+,\d{2})"
)


def _dividir_linhas_parcela_concatenadas(linhas_pagina):
    """Quando ha mais de uma parcela prevista de liberacao, o PRN as
    vezes concatena todas na MESMA linha fisica (ex: '1 DATA VALOR PCT
    2 DATA VALOR PCT' tudo junto, sem quebra de linha entre elas) --
    detecta isso e separa cada parcela na sua propria linha, no mesmo
    formato de coluna da primeira."""
    resultado = []
    for linha in linhas_pagina:
        ocorrencias = list(PADRAO_LINHA_PARCELA.finditer(linha))
        if len(ocorrencias) >= 2:
            for m in ocorrencias:
                nro, data, valor, pct = m.groups()
                resultado.append(f"        {nro}   {data}            {valor}           {pct}")
        else:
            resultado.append(linha)
    return resultado


def _reflow_prosa(linhas_texto):
    """Junta linhas quebradas por largura fixa em texto corrido, normalizando
    espacamento duplo (usado no PRN para justificar) para espaco simples."""
    texto = " ".join(l.strip() for l in linhas_texto if l.strip())
    return re.sub(r" {2,}", " ", texto)


def _classificar_bloco(linhas_texto):
    if any(PADRAO_UNDERLINE.match(l) for l in linhas_texto):
        return "assinatura"

    linhas_stripped = [l.strip() for l in linhas_texto]
    # so classifica como tabular com marcadores confiaveis: linha
    # separadora (----/====) ou cabecalho de tabela conhecido. A
    # heuristica antiga (maioria das linhas com espaco largo) foi
    # removida -- texto justificado normal tambem pode ter gaps largos
    # quando a linha tem poucas palavras pra preencher a largura, e
    # isso gerava falso positivo em clausulas normais.
    if any(PADRAO_SEPARADOR.match(l) or PADRAO_CABECALHO_TABELA.search(l) for l in linhas_stripped):
        return "tabular"

    if any(PADRAO_CAMPO_FORMULARIO.search(l) for l in linhas_texto):
        return "formulario"
    return "prosa"


def _adicionar_campo_pagina(paragraph):
    """Insere um campo de numero de pagina automatico (se atualiza sozinho
    no Word, nao e texto fixo)."""
    run = paragraph.add_run()
    fld_inicio = OxmlElement("w:fldChar")
    fld_inicio.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_fim = OxmlElement("w:fldChar")
    fld_fim.set(qn("w:fldCharType"), "end")
    run._r.append(fld_inicio)
    run._r.append(instr)
    run._r.append(fld_fim)
    run.font.name = "Courier New"
    run.font.size = Pt(FONTE_PT)
    run.font.color.rgb = RGBColor(0, 0, 0)


def _adicionar_texto_exceto_ultima_pagina(paragraph, texto_condicional):
    """Insere um campo condicional do Word: so mostra 'texto_condicional'
    quando a pagina atual NAO for a ultima pagina da secao (campo IF com
    PAGE e SECTIONPAGES aninhados). Assim o rodape "Continua Proxima
    Pagina" desaparece sozinho na ultima pagina, sem precisar criar
    nenhuma secao/pagina extra -- o Word calcula isso dinamicamente,
    entao continua correto mesmo se o texto for editado depois."""

    def _fld_char(tipo):
        el = OxmlElement("w:fldChar")
        el.set(qn("w:fldCharType"), tipo)
        return el

    def _instr(texto):
        el = OxmlElement("w:instrText")
        el.set(qn("xml:space"), "preserve")
        el.text = texto
        return el

    run = paragraph.add_run()
    run.font.name = "Courier New"
    run.font.size = Pt(FONTE_PT)
    run.font.color.rgb = RGBColor(0, 0, 0)
    r = run._r

    r.append(_fld_char("begin"))  # campo IF (externo)
    r.append(_instr(' IF '))
    r.append(_fld_char("begin"))  # campo PAGE (aninhado)
    r.append(_instr(' PAGE '))
    r.append(_fld_char("end"))
    r.append(_instr(' <> '))
    r.append(_fld_char("begin"))  # campo SECTIONPAGES (aninhado)
    r.append(_instr(' SECTIONPAGES '))
    r.append(_fld_char("end"))
    r.append(_instr(f' "{texto_condicional}" ""'))
    r.append(_fld_char("separate"))
    # texto de exibicao inicial (o Word recalcula ao abrir o arquivo)
    texto_inicial = OxmlElement("w:t")
    texto_inicial.set(qn("xml:space"), "preserve")
    texto_inicial.text = texto_condicional
    r.append(texto_inicial)
    r.append(_fld_char("end"))


PADRAO_CARTILHA = re.compile(r"CARTILHA\s+DO\s+CR[EÉ]DITO\s+RURAL:", re.IGNORECASE)


PADRAO_ITEM_LISTA_CARTILHA = re.compile(r"^\s*-\s")
PADRAO_TERMINA_FRASE = re.compile(r"[.:;]\s*$")
PADRAO_SICREDI_FONE = re.compile(r"SICREDI\s+FONE", re.IGNORECASE)


def _mesclar_topicos_cartilha(linhas):
    """Junta linhas quebradas por largura fixa em um paragrafo por
    'topico' (paragrafo, cabecalho ou item de lista com '-'), evitando
    que um heading fique despedacado por uma quebra de pagina do PRN
    original caindo no meio da frase (2+ linhas em branco seguidas e a
    frase anterior nao termina em pontuacao = artefato de pagina, nao
    paragrafo novo). Itens de lista com '-' sempre comecam um topico
    novo, mesmo sem linha em branco antes (e como vem no PRN)."""
    resultado = []
    topico_atual = []

    def ultima_linha_nao_vazia():
        for l in reversed(topico_atual):
            if l.strip():
                return l.strip()
        return ""

    def fechar_topico():
        if topico_atual:
            resultado.append(_reflow_prosa(topico_atual))
            topico_atual.clear()

    i = 0
    while i < len(linhas):
        linha = linhas[i]
        if linha.strip() == "":
            j = i
            while j < len(linhas) and linhas[j].strip() == "":
                j += 1
            n_blanks = j - i
            ultima = ultima_linha_nao_vazia()
            termina_frase = bool(PADRAO_TERMINA_FRASE.search(ultima)) if ultima else True

            if n_blanks == 1 or termina_frase:
                fechar_topico()
                resultado.append("")
            # senao (2+ linhas em branco e a frase anterior nao termina em
            # pontuacao): artefato de quebra de pagina no meio da frase --
            # ignora as linhas em branco e continua o topico atual
            i = j
            continue

        if PADRAO_ITEM_LISTA_CARTILHA.match(linha) and topico_atual:
            fechar_topico()

        topico_atual.append(linha)
        i += 1

    fechar_topico()
    return resultado


def gerar_docx(texto, blocos, assinaturas_qualificadas, caminho_saida):
    # a partir da "CARTILHA DO CREDITO RURAL:" (se houver), o documento
    # nao usa mais a paginacao "Continuacao do instrumento..." no PRN
    # original -- e um anexo informativo separado. Por pedido do usuario,
    # essa parte fica fiel ao original (sem reflow/tabela/qualificacao,
    # sem cabecalho/rodape fixo), como uma segunda secao do Word.
    m_cartilha = PADRAO_CARTILHA.search(texto)
    if m_cartilha:
        texto_estruturado = texto[: m_cartilha.start()]
        texto_verbatim = texto[m_cartilha.start():]
    else:
        texto_estruturado = texto
        texto_verbatim = ""

    paginas_textos = texto_estruturado.split("\x0c")

    # posicao (no texto todo) do ultimo campo de assinatura -- so
    # removemos paginas em branco QUE VENHAM DEPOIS desse ponto (efeito
    # colateral da reestruturacao do Interveniente Garantidor, que pode
    # deixar uma pagina so com "Continua Proxima Pagina").
    pos_ultima_assinatura = 0
    for m in re.finditer(r"\bNOME:\s*.+", texto_estruturado, re.IGNORECASE):
        pos_ultima_assinatura = m.end()

    paginas_linhas = []
    offsets_paginas = []
    offset = 0
    for pagina in paginas_textos:
        offsets_paginas.append(offset)
        offset += len(pagina) + 1  # +1 pelo \x0c removido no split
        linhas_pagina = [l.replace("\r", "") for l in pagina.replace("\r\n", "\n").split("\n")]
        linhas_pagina = _dividir_linhas_parcela_concatenadas(linhas_pagina)
        paginas_linhas.append(linhas_pagina)

    while paginas_linhas and all(l.strip() == "" for l in paginas_linhas[-1]):
        paginas_linhas.pop()
        offsets_paginas.pop()

    i = 1  # nunca remove a pagina 0 (cabecalho do titulo)
    while i < len(paginas_linhas):
        depois_da_ultima_assinatura = offsets_paginas[i] >= pos_ultima_assinatura
        if depois_da_ultima_assinatura and _pagina_so_tem_boilerplate(paginas_linhas[i]):
            paginas_linhas[i - 1] = [l for l in paginas_linhas[i - 1] if not PADRAO_CONTINUA_PROXIMA.search(l)]
            paginas_linhas.pop(i)
            offsets_paginas.pop(i)
            # nao incrementa i -- a lista encolheu
        else:
            i += 1

    for linhas in paginas_linhas:
        idx = 0
        while idx < len(linhas) - 1:
            if PADRAO_AUTORIZACAO_L1.search(linhas[idx]) and PADRAO_AUTORIZACAO_L2.search(linhas[idx + 1]):
                linhas[idx] = linhas[idx].rstrip() + " " + linhas[idx + 1].strip()
                linhas.pop(idx + 1)
            idx += 1

    apartir_sicor = False
    linhas_flat = []
    for idx_pagina, linhas in enumerate(paginas_linhas):
        for idx_linha, linha in enumerate(linhas):
            if MARCADOR_FONTE_MENOR in linha:
                apartir_sicor = True
            linhas_flat.append({
                "texto": linha,
                # NAO forca quebra de pagina nos limites originais do PRN --
                # agora que o texto reflui em prosa compacta, isso so
                # desperdicava espaco (paginas terminando com muito branco
                # sobrando). O conteudo flui naturalmente; so o bloco de
                # assinatura continua protegido contra ficar dividido.
                "quebra_forcada": False,
                "fonte_menor": apartir_sicor,
            })

    indice_assinatura = 0
    unidades = []
    i = 0
    while i < len(linhas_flat):
        atual = linhas_flat[i]

        eh_marcador_paginacao = (
            PADRAO_CONTINUA_PROXIMA.search(atual["texto"])
            or PADRAO_CONTINUACAO_HEADER.search(atual["texto"])
            or PADRAO_PAGINA_SOZINHA.match(atual["texto"])
        )
        if eh_marcador_paginacao:
            i += 1
            continue  # nao vira paragrafo -- agora vem do cabecalho/rodape fixo
        if atual["texto"].strip() == "":
            unidades.append({"tipo": "linha", "linha": atual})
            i += 1
            continue

        # junta linhas nao-vazias consecutivas em um bloco, pra classificar
        # (para tambem nos marcadores de paginacao, que ficam sempre
        # isolados, E quando o bloco comeca com um sublinhado de
        # assinatura, tambem para ao achar OUTRO sublinhado -- senao,
        # quando duas assinaturas vem coladas sem linha em branco entre
        # elas, a segunda fica perdida dentro do bloco da primeira)
        eh_bloco_assinatura = bool(PADRAO_UNDERLINE.match(atual["texto"]))
        bloco = [atual]
        j = i + 1
        while j < len(linhas_flat) and linhas_flat[j]["texto"].strip() != "" and not (
            PADRAO_CONTINUA_PROXIMA.search(linhas_flat[j]["texto"])
            or PADRAO_CONTINUACAO_HEADER.search(linhas_flat[j]["texto"])
            or PADRAO_PAGINA_SOZINHA.match(linhas_flat[j]["texto"])
            or (eh_bloco_assinatura and PADRAO_UNDERLINE.match(linhas_flat[j]["texto"]))
        ):
            bloco.append(linhas_flat[j])
            j += 1

        tipo = _classificar_bloco([l["texto"] for l in bloco])

        if tipo == "assinatura":
            linhas_bloco = [{**bloco[0], "texto": bloco[0]["texto"].strip()}]
            resto = bloco[1:]
            if resto and PADRAO_AUTORIZACAO_UMA_LINHA.search(resto[0]["texto"]):
                linhas_bloco.append({**resto[0], "texto": resto[0]["texto"].strip()})
                resto = resto[1:]
            # rotulo solto "Interveniente(s) Garantidor(es)" (sem ":"),
            # que pode aparecer sozinho dentro de outro bloco (ex: dentro
            # do EMITENTE(S)) -- mantem visivel e avanca, senao a leitura
            # sequencial trava aqui e o Nome/CPF seguintes ficam perdidos
            # (e todas as assinaturas depois saem deslocadas/erradas).
            if resto and PADRAO_INTERVENIENTE_LABEL_SOLTO.match(resto[0]["texto"].strip()):
                linhas_bloco.append({**resto[0], "texto": resto[0]["texto"].strip()})
                resto = resto[1:]
            # a partir daqui, busca as proximas linhas "de verdade" pulando
            # linhas em branco e marcadores de paginacao (a assinatura pode
            # atravessar quebra de pagina -- ex: "Razao Social:" numa
            # pagina e o "CNPJ" na seguinte). Para se achar OUTRO sublinhado.
            cursor = i + 1
            linha, cursor = _proxima_linha_assinatura(linhas_flat, cursor)

            if linha and PADRAO_AUTORIZACAO_UMA_LINHA.search(linha["texto"]):
                linhas_bloco.append({**linha, "texto": linha["texto"].strip()})
                cursor += 1
                linha, cursor = _proxima_linha_assinatura(linhas_flat, cursor)

            if linha and PADRAO_INTERVENIENTE_LABEL_SOLTO.match(linha["texto"].strip()):
                linhas_bloco.append({**linha, "texto": linha["texto"].strip()})
                cursor += 1
                linha, cursor = _proxima_linha_assinatura(linhas_flat, cursor)

            eh_pj_aqui = linha is not None and PADRAO_RAZAO_SOCIAL.search(linha["texto"])
            eh_pf_aqui = linha is not None and re.search(r"\bNOME:\s*.+", linha["texto"], re.IGNORECASE)

            if linha and (eh_pj_aqui or eh_pf_aqui):
                linha_nome = linha
                linha_doc, cursor_doc = _proxima_linha_assinatura(linhas_flat, cursor + 1)
                if eh_pj_aqui:
                    tem_doc = linha_doc is not None and re.search(r"CNPJ", linha_doc["texto"], re.IGNORECASE)
                else:
                    tem_doc = linha_doc is not None and re.match(r"^\s*CPF", linha_doc["texto"], re.IGNORECASE)

                assinatura = assinaturas_qualificadas[indice_assinatura] if indice_assinatura < len(assinaturas_qualificadas) else None

                if assinatura and tem_doc:
                    indice_assinatura += 1
                    dados = assinatura["dados_completos"]
                    bloco_texto = dados.get("bloco_sem_rotulo") if dados else None
                    if bloco_texto:
                        linhas_originais = bloco_texto.split("\r\n") if "\r\n" in bloco_texto else bloco_texto.split("\n")
                        texto_reflow = _reflow_prosa(linhas_originais)
                        linhas_bloco.append({"texto": texto_reflow, "quebra_forcada": False, "fonte_menor": atual["fonte_menor"]})
                        # PJ: acrescenta o modelo de representacao por socio
                        # (com asteriscos pro colaborador preencher -- o PRN
                        # nao traz nome/CPF do socio), marcado em azul
                        if assinatura.get("eh_pj"):
                            linhas_bloco.append({
                                "texto": MARCADOR_ADICIONADO_INICIO + TEXTO_REPRESENTACAO_PJ + MARCADOR_ADICIONADO_FIM,
                                "quebra_forcada": False, "fonte_menor": atual["fonte_menor"],
                            })
                    else:
                        linhas_bloco.append(linha_nome)
                        if linha_doc:
                            linhas_bloco.append(linha_doc)
                    j = cursor_doc + 1 if tem_doc else cursor + 1
                else:
                    linhas_bloco.append(linha_nome)
                    if tem_doc:
                        linhas_bloco.append(linha_doc)
                    j = cursor_doc + 1 if tem_doc else cursor + 1
            else:
                # nao achou Nome/Razao Social (proximo sublinhado ou fim do
                # documento) -- retoma dali, sem consumir nada a mais
                j = cursor

            unidades.append({"tipo": "bloco", "linhas": linhas_bloco, "quebra_forcada": atual["quebra_forcada"]})

        elif tipo == "tabular":
            for l in bloco:
                unidades.append({"tipo": "linha", "linha": l})

        elif tipo == "formulario":
            # mantem cada linha separada (campo a campo), so normaliza
            # espacamento duplo e tira o recuo fixo
            for l in bloco:
                l2 = {**l, "texto": re.sub(r" {2,}", " ", l["texto"].strip())}
                unidades.append({"tipo": "linha", "linha": l2})

        else:  # prosa
            texto_reflow = _reflow_prosa([l["texto"] for l in bloco])
            unidades.append({
                "tipo": "linha",
                "linha": {"texto": texto_reflow, "quebra_forcada": atual["quebra_forcada"], "fonte_menor": atual["fonte_menor"]},
            })
            j = j  # (mantem o j calculado no agrupamento inicial do bloco)

        i = j

    doc = Document()
    secao = doc.sections[0]
    secao.orientation = WD_ORIENT.PORTRAIT
    secao.page_width = Cm(21.59)   # 8.5in em retrato = largura (Oficio/Legal)
    secao.page_height = Cm(35.56)  # 14in em retrato = altura
    # topo com folga suficiente pra caber o cabecalho (distancia + altura da
    # linha) sem invadir o espaco do corpo do texto -- se nao, o corpo
    # "perde" espaco que o calculo de linhas por pagina nao contava, e
    # sobra um pouco de texto vazando pra pagina seguinte.
    secao.top_margin = Cm(0.7)
    secao.bottom_margin = Cm(1.25)
    secao.left_margin = Cm(1.5)
    secao.right_margin = Cm(1.5)

    # -------- cabecalho e rodape fixos (Continuacao/Pagina/Continua Proxima) --------
    m_titulo = re.search(r"T[IÍ]TULO\.+:\s*(\S+)", texto, re.IGNORECASE)
    numero_titulo = m_titulo.group(1) if m_titulo else ""
    largura_util_cm = 21.59 - 1.5 - 1.5  # largura da pagina menos as margens laterais

    secao.header_distance = Cm(0.15)
    secao.footer_distance = Cm(0.3)
    secao.different_first_page_header_footer = True

    PRETO = RGBColor(0, 0, 0)

    def _run_padrao(paragraph, texto_run):
        r = paragraph.add_run(texto_run)
        r.font.name = "Courier New"
        r.font.size = Pt(FONTE_PT)
        r.font.color.rgb = PRETO
        return r

    # cabecalho da 1a pagina: so "Pagina: 1", alinhado a direita
    p_header_1 = secao.first_page_header.paragraphs[0]
    p_header_1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _run_padrao(p_header_1, "Pagina: ")
    _adicionar_campo_pagina(p_header_1)

    # cabecalho das demais paginas: "Continuacao do instrumento..." a
    # esquerda + "Pagina: N" (campo automatico) a direita, usando tab
    # com parada alinhada a direita na borda util da pagina
    p_header = secao.header.paragraphs[0]
    p_header.paragraph_format.tab_stops.add_tab_stop(Cm(largura_util_cm), WD_TAB_ALIGNMENT.RIGHT)
    _run_padrao(p_header, f"Continuação do instrumento de crédito do título {numero_titulo}.\tPagina: ")
    _adicionar_campo_pagina(p_header)

    # rodape (todas as paginas, exceto a ultima da secao): "Continua
    # Proxima Pagina" alinhado a direita -- campo condicional, some
    # sozinho na ultima pagina (onde termina o instrumento e comeca a
    # Cartilha, um documento diferente).
    for rodape in (secao.footer, secao.first_page_footer):
        p_footer = rodape.paragraphs[0]
        p_footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        _adicionar_texto_exceto_ultima_pagina(p_footer, "Continua Proxima Pagina")
    # -------------------------------------------------------------------------

    altura_linha_pt = FONTE_PT * FATOR_ALTURA_LINHA

    def adicionar_paragrafo(texto_linha, fonte_menor, manter_com_proximo=False, manter_junto=False):
        eh_adicionado = MARCADOR_ADICIONADO_INICIO in texto_linha
        texto_limpo = texto_linha.replace(MARCADOR_ADICIONADO_INICIO, "").replace(MARCADOR_ADICIONADO_FIM, "")

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pf = p.paragraph_format
        pf.space_after = Pt(0)
        pf.space_before = Pt(0)
        # AT_LEAST (nao EXACTLY): se uma linha longa quebrar em 2 linhas
        # visuais (retrato e mais estreito que paisagem), o paragrafo
        # cresce em vez de sobrepor o texto da linha seguinte.
        pf.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
        pf.line_spacing = Pt(FONTE_MENOR_PT * FATOR_ALTURA_LINHA if fonte_menor else altura_linha_pt)
        # deixa o WORD decidir onde quebrar a pagina (dinamico de
        # verdade) -- so pedimos pra ele manter certos paragrafos juntos,
        # em vez de calcular por fora quantas linhas cabem.
        pf.keep_with_next = manter_com_proximo
        pf.keep_together = manter_junto
        run = p.add_run(texto_limpo if len(texto_limpo) else " ")
        run.font.name = "Courier New"
        run.font.size = Pt(FONTE_MENOR_PT if fonte_menor else FONTE_PT)
        if eh_adicionado:
            run.font.color.rgb = RGBColor(0, 51, 153)  # azul -- sinaliza texto acrescentado pelo script
        return p

    for u in unidades:
        if u["tipo"] == "linha":
            adicionar_paragrafo(u["linha"]["texto"], u["linha"]["fonte_menor"])
        else:
            n_linhas_bloco = len(u["linhas"])
            for idx, l in enumerate(u["linhas"]):
                eh_ultima_linha_do_bloco = idx == n_linhas_bloco - 1
                adicionar_paragrafo(
                    l["texto"], l["fonte_menor"],
                    manter_com_proximo=not eh_ultima_linha_do_bloco,
                    manter_junto=True,
                )

    # -------- 2a e 3a secoes: Cartilha e Extrato SICOR, fieis ao original -----
    # nenhuma das regras (reflow/tabela/qualificacao/cabecalho fixo) se
    # aplica aqui -- cada linha do PRN vira um paragrafo tal como veio,
    # e essas secoes nao tem cabecalho/rodape (a paginacao "oficial" do
    # instrumento acaba na ultima assinatura). Margens quase zero nas
    # duas ("sem espaçamento"); a do Extrato SICOR fica em fonte 7pt.
    m_sicor = re.search(re.escape(MARCADOR_FONTE_MENOR), texto_verbatim)
    if m_sicor:
        texto_cartilha_only = texto_verbatim[: m_sicor.start()]
        texto_sicor_only = texto_verbatim[m_sicor.start():]
    else:
        texto_cartilha_only = texto_verbatim
        texto_sicor_only = ""

    def _nova_secao_sem_espacamento(margem_lateral):
        s = doc.add_section(WD_SECTION.NEW_PAGE)
        s.page_width = secao.page_width
        s.page_height = secao.page_height
        s.top_margin = Cm(0.2)
        s.bottom_margin = Cm(0.2)
        s.left_margin = margem_lateral
        s.right_margin = margem_lateral
        s.header.is_linked_to_previous = False
        s.footer.is_linked_to_previous = False
        s.first_page_header.is_linked_to_previous = False
        s.first_page_footer.is_linked_to_previous = False
        for p in s.header.paragraphs:
            p.text = ""
        for p in s.footer.paragraphs:
            p.text = ""
        for p in s.first_page_header.paragraphs:
            p.text = ""
        for p in s.first_page_footer.paragraphs:
            p.text = ""
        return s

    def _adicionar_paginas_verbatim(texto_bloco, tamanho_fonte, mesclar_topicos=False):
        linhas_todas = []
        for pagina in texto_bloco.split("\x0c"):
            linhas_todas.extend(l.replace("\r", "") for l in pagina.replace("\r\n", "\n").split("\n"))

        linhas_centralizadas = set()

        if mesclar_topicos:
            # separa o bloco final "SICREDI FONE" (4 linhas fixas de
            # contato) do resto -- ele NAO deve ser mesclado com o
            # topico anterior, e cada uma das 4 linhas fica centralizada
            # e sem o preenchimento de espacos antigo (calculado pra
            # largura de pagina antiga).
            idx_sicredi = next(
                (i for i, l in enumerate(linhas_todas) if PADRAO_SICREDI_FONE.search(l)), None
            )
            if idx_sicredi is not None:
                antes = linhas_todas[:idx_sicredi]
                bloco_contato = [l.strip() for l in linhas_todas[idx_sicredi:idx_sicredi + 4] if l.strip()]
                depois = linhas_todas[idx_sicredi + 4:]

                topicos = _mesclar_topicos_cartilha(antes)
                # tira blanks do final e poe exatamente 3 antes do bloco de contato
                while topicos and topicos[-1] == "":
                    topicos.pop()
                topicos.extend([""] * 3)

                inicio_centralizadas = len(topicos)
                topicos.extend(bloco_contato)
                linhas_centralizadas = set(range(inicio_centralizadas, len(topicos)))

                topicos.extend(_mesclar_topicos_cartilha(depois))
                linhas_todas = topicos
            else:
                linhas_todas = _mesclar_topicos_cartilha(linhas_todas)

        altura_linha_local_pt = tamanho_fonte * FATOR_ALTURA_LINHA

        # deixa o Word decidir onde quebrar a pagina -- so pede pra
        # manter cada topico/linha inteiro junto (nao dividir no meio),
        # sem calcular quantas linhas cabem por pagina.
        for idx_linha, linha in enumerate(linhas_todas):
            p = doc.add_paragraph()
            centralizar = idx_linha in linhas_centralizadas or MARCADOR_FONTE_MENOR in linha
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if centralizar else WD_ALIGN_PARAGRAPH.LEFT
            pf = p.paragraph_format
            pf.space_after = Pt(0)
            pf.space_before = Pt(0)
            pf.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
            pf.line_spacing = Pt(altura_linha_local_pt)
            pf.keep_together = True
            run = p.add_run(linha if linha else " ")
            run.font.name = "Courier New"
            run.font.size = Pt(tamanho_fonte)
            run.font.color.rgb = RGBColor(0, 0, 0)

    if texto_cartilha_only:
        _nova_secao_sem_espacamento(Cm(1.5))
        _adicionar_paginas_verbatim(texto_cartilha_only, FONTE_PT, mesclar_topicos=True)

    if texto_sicor_only:
        _nova_secao_sem_espacamento(Cm(0))
        _adicionar_paginas_verbatim(texto_sicor_only, 7)

    doc.save(caminho_saida)


# ======================================================================
# ORQUESTRACAO
# ======================================================================

def sanitizar_nome_arquivo(s):
    return re.sub(r'[\\/:*?"<>|]', "-", s).strip()


def preparar_arquivo(caminho_prn, log=print):
    """Decodifica, reestrutura, qualifica e detecta a secao de imoveis.
    NAO gera o Word ainda -- devolve tudo que a tela de revisao precisa."""
    log(f"Decodificando {os.path.basename(caminho_prn)} ...")
    texto = decodificar_prn(caminho_prn)
    texto = reestruturar_documento(texto)

    log("Extraindo blocos de dados e qualificando assinaturas ...")
    blocos = extrair_blocos_dados(texto)
    assinaturas = extrair_assinaturas(texto)
    assinaturas_qualificadas = qualificar_com_blocos(assinaturas, blocos)
    log(f"  {len(assinaturas_qualificadas)} assinatura(s) qualificada(s).")

    deteccao_imoveis = detectar_secao_imoveis(texto)
    if deteccao_imoveis:
        avisos = []
        if deteccao_imoveis["tem_hipoteca"]:
            avisos.append("hipoteca encontrada")
        if deteccao_imoveis["avaliacoes_duplicadas"]:
            avisos.append("valores de avaliação idênticos (possível duplicidade)")
        if avisos:
            log("  Seção '2 - IMÓVEIS': " + "; ".join(avisos) + ".")
        else:
            log("  Seção '2 - IMÓVEIS' encontrada (sem hipoteca, valores de avaliação distintos).")

    tem_clausula_superveniencia = detectar_clausula_superveniencia(texto)
    cartorio_detectado = detectar_cartorio(texto)
    log(f"  Cláusula de Superveniência: {'encontrada' if tem_clausula_superveniencia else 'não encontrada'}.")
    log(f"  Cartório detectado: {cartorio_detectado or '(nenhum)'}")

    deteccao_hipoteca = detectar_secao_hipoteca(texto)
    if deteccao_hipoteca:
        if deteccao_hipoteca["tinha_duplicacao"]:
            log("  Hipoteca Cedular encontrada -- trecho duplicado detectado e será corrigido.")
        else:
            log("  Hipoteca Cedular encontrada (sem duplicação).")

    emitente = next((b for b in blocos if b["papel"] == "EMITENTE"), None)
    nome_emitente = emitente["nome"].strip() if emitente and emitente.get("nome") else None
    m_titulo = re.search(r"T[IÍ]TULO\.+:\s*(\S+)", texto, re.IGNORECASE)
    numero_titulo = m_titulo.group(1) if m_titulo else None

    return {
        "caminho_prn": caminho_prn,
        "texto": texto,
        "blocos": blocos,
        "assinaturas_qualificadas": assinaturas_qualificadas,
        "deteccao_imoveis": deteccao_imoveis,
        "tem_clausula_superveniencia": tem_clausula_superveniencia,
        "cartorio_detectado": cartorio_detectado,
        "deteccao_hipoteca": deteccao_hipoteca,
        "nome_emitente": nome_emitente,
        "numero_titulo": numero_titulo,
    }


def finalizar_geracao(
    dados, pasta_saida, texto_revisao_imoveis=None, checkbox_superveniencia=None,
    texto_revisao_assinaturas=None, agencia_procurador=None, log=print,
):
    """Aplica a correcao da secao de imoveis (se houver texto revisado),
    acrescenta a clausula de Superveniencia se o checkbox estiver marcado
    E ela nao existia no documento original (evita duplicidade), corrige
    a Hipoteca Cedular se houver duplicacao, aplica a revisao manual das
    assinaturas (se editada), acrescenta a assinatura da CREDORA
    representada pelos procuradores da agencia escolhida (se houver), e
    gera o Word."""
    texto = dados["texto"]
    posicao_fim_descricao = dados["deteccao_imoveis"]["span_fim"] if dados["deteccao_imoveis"] else None

    if dados["deteccao_imoveis"] and texto_revisao_imoveis is not None:
        log("Aplicando correções da seção '2 - IMÓVEIS' ...")
        texto, posicao_fim_descricao = aplicar_correcao_imoveis(texto, dados["deteccao_imoveis"], texto_revisao_imoveis)

    # 3 casos pro checkbox de Superveniencia:
    #  - ja existia e continua marcado -> nao faz nada (evita duplicidade)
    #  - ja existia e foi DESMARCADO -> remove a clausula do documento
    #  - nao existia e foi marcado -> acrescenta (marcada em azul)
    deve_acrescentar_clausula = (
        checkbox_superveniencia
        and not dados["tem_clausula_superveniencia"]
        and posicao_fim_descricao is not None
    )
    deve_remover_clausula = (
        checkbox_superveniencia is False
        and dados["tem_clausula_superveniencia"]
    )
    if deve_acrescentar_clausula:
        log("Acrescentando cláusula de Superveniência (marcada em azul no Word) ...")
        texto = aplicar_clausula_superveniencia(texto, posicao_fim_descricao, adicionar=True)
    elif deve_remover_clausula:
        log("Removendo cláusula de Superveniência (desmarcada pelo colaborador) ...")
        texto = remover_clausula_superveniencia(texto)

    # Hipoteca Cedular: corrige duplicacao automaticamente, se houver.
    # Redeteta em cima do texto ATUAL (nao o original) pra nao usar
    # posicoes desatualizadas depois das edicoes acima.
    deteccao_hipoteca_atual = detectar_secao_hipoteca(texto)
    hipoteca_corrigida = deteccao_hipoteca_atual is not None and deteccao_hipoteca_atual["tinha_duplicacao"]
    if hipoteca_corrigida:
        log("Corrigindo trecho duplicado na cláusula de Hipoteca Cedular ...")
        texto = aplicar_correcao_hipoteca(texto, deteccao_hipoteca_atual)

    texto_mudou = (
        texto_revisao_imoveis is not None
        or deve_acrescentar_clausula
        or deve_remover_clausula
        or hipoteca_corrigida
    )
    if texto_mudou:
        # o texto mudou -- refaz blocos/assinaturas em cima do texto
        # corrigido, pra qualificacao continuar correta.
        blocos = extrair_blocos_dados(texto)
        assinaturas_qualificadas = qualificar_com_blocos(extrair_assinaturas(texto), blocos)
    else:
        blocos = dados["blocos"]
        assinaturas_qualificadas = dados["assinaturas_qualificadas"]

    if texto_revisao_assinaturas is not None:
        log("Aplicando revisão das assinaturas ...")
        texto = aplicar_revisao_assinaturas(texto, assinaturas_qualificadas, texto_revisao_assinaturas, log=log)

    if agencia_procurador:
        log(f"Acrescentando assinatura da CREDORA (procuradores da agência {agencia_procurador['rotulo']}) ...")
        texto = aplicar_assinatura_procurador(texto, agencia_procurador)

    if dados["nome_emitente"] and dados["numero_titulo"]:
        nome_arquivo = sanitizar_nome_arquivo(f"{dados['nome_emitente']} - {dados['numero_titulo']}") + ".docx"
    else:
        nome_arquivo = os.path.splitext(os.path.basename(dados["caminho_prn"]))[0] + "_decodificado.docx"

    caminho_saida = os.path.join(pasta_saida, nome_arquivo)

    log("Gerando o Word ...")
    try:
        gerar_docx(texto, blocos, assinaturas_qualificadas, caminho_saida)
    except PermissionError:
        alternativo = os.path.join(pasta_saida, f"{os.path.splitext(nome_arquivo)[0]}_novo.docx")
        log(f'AVISO: "{nome_arquivo}" esta aberto/bloqueado. Salvando como {os.path.basename(alternativo)} ...')
        gerar_docx(texto, blocos, assinaturas_qualificadas, alternativo)
        caminho_saida = alternativo

    log(f"Pronto: {caminho_saida}")
    return caminho_saida


# ======================================================================
# INTERFACE GRAFICA
# ======================================================================

class App(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title("Qualificador de Cédulas")
        self.geometry("760x620")
        self.resizable(True, True)

        self.caminho_prn = tk.StringVar()
        self.var_superveniencia = tk.BooleanVar(value=False)
        self.var_cartorio = tk.StringVar()
        self.var_procurador = tk.StringVar()
        self.dados_preparados = None  # resultado de preparar_arquivo()
        self.agencias_procuradores = []  # lista de dicts vinda da planilha
        pasta_documentos()  # garante que a pasta exista, pronta pra planilha ser colocada

        frame_top = tk.Frame(self, padx=12, pady=12)
        frame_top.pack(fill="x")
        tk.Label(frame_top, text="Arquivo .PRN:").pack(side="left")
        tk.Entry(frame_top, textvariable=self.caminho_prn, width=55).pack(side="left", padx=6, fill="x", expand=True)
        tk.Button(frame_top, text="Selecionar...", command=self.selecionar_arquivo).pack(side="left")

        frame_botoes = tk.Frame(self, padx=12)
        frame_botoes.pack(fill="x")
        self.botao_processar = tk.Button(
            frame_botoes, text="1. Processar", command=self.processar, bg="#1565c0", fg="white"
        )
        self.botao_processar.pack(side="left", pady=6)
        self.botao_gerar = tk.Button(
            frame_botoes, text="2. Gerar Word", command=self.gerar, bg="#2e7d32", fg="white", state="disabled"
        )
        self.botao_gerar.pack(side="left", padx=(8, 0), pady=6)

        # -------- abas --------
        self.abas = ttk.Notebook(self)
        self.abas.pack(fill="both", expand=True, padx=12, pady=(6, 6))

        self._montar_aba_descricao()
        self._montar_aba_assinaturas()
        self._montar_aba_opcoes()

        self.log_text = scrolledtext.ScrolledText(self, padx=8, pady=8, state="disabled", wrap="word", height=8)
        self.log_text.pack(fill="both", expand=False, padx=12, pady=(0, 12))

    # ------------------------------------------------------------------
    # montagem das abas
    # ------------------------------------------------------------------
    def _montar_aba_descricao(self):
        aba = tk.Frame(self.abas, padx=8, pady=8)
        self.abas.add(aba, text="Descrição da Matrícula")

        self.aviso_label = tk.Label(aba, text="", fg="#b71c1c", justify="left", anchor="w")
        self.aviso_label.pack(fill="x", anchor="w")

        tk.Label(
            aba, text="Descrição do imóvel (seção \"2 - IMÓVEIS\") -- edite aqui se precisar corrigir:",
            anchor="w", justify="left",
        ).pack(fill="x", anchor="w", pady=(6, 2))

        self.texto_revisao = scrolledtext.ScrolledText(aba, height=16, wrap="word")
        self.texto_revisao.pack(fill="both", expand=True)

    def _montar_aba_assinaturas(self):
        aba = tk.Frame(self.abas, padx=8, pady=8)
        self.abas.add(aba, text="Assinaturas")

        tk.Label(
            aba,
            text="Um bloco por assinatura (emitente/avalista/PJ), com todos os dados que vão pro\n"
                 "Word. Edite pra corrigir algo, apague o texto de um bloco pra remover a\n"
                 "qualificação dele, ou copie o formato \"--- N. PAPEL (documento) ---\" pra\n"
                 "acrescentar uma assinatura nova (numere maior que as existentes).",
            fg="#555555", justify="left", anchor="w",
        ).pack(fill="x", anchor="w", pady=(0, 4))

        self.texto_assinaturas = scrolledtext.ScrolledText(aba, height=16, wrap="word")
        self.texto_assinaturas.pack(fill="both", expand=True)

    def _montar_aba_opcoes(self):
        aba = tk.Frame(self.abas, padx=12, pady=12)
        self.abas.add(aba, text="Opções")

        frame_superveniencia = tk.Frame(aba)
        frame_superveniencia.pack(fill="x", anchor="w", pady=(0, 4))
        self.check_superveniencia = tk.Checkbutton(
            frame_superveniencia, text="Cláusula de Superveniência", variable=self.var_superveniencia,
        )
        self.check_superveniencia.pack(side="left")
        self.label_status_superveniencia = tk.Label(frame_superveniencia, text="", fg="#555555")
        self.label_status_superveniencia.pack(side="left", padx=(8, 0))

        tk.Label(
            aba,
            text="Se veio desmarcado, a cláusula não foi encontrada -- marque para acrescentá-la\n"
                 "(aparece em azul no Word). Se veio marcado (já existia) e você desmarcar, ela é\n"
                 "removida do documento. Deixar marcado como veio não duplica.",
            fg="#555555", justify="left", anchor="w",
        ).pack(fill="x", anchor="w", pady=(0, 12))

        tk.Label(aba, text="Cartório:").pack(anchor="w")
        self.combo_cartorio = ttk.Combobox(aba, textvariable=self.var_cartorio, values=[], width=60)
        self.combo_cartorio.pack(fill="x", anchor="w", pady=(0, 12))

        tk.Label(aba, text="Procuradores:").pack(anchor="w")
        self.combo_procurador = ttk.Combobox(aba, textvariable=self.var_procurador, values=[], width=60)
        self.combo_procurador.pack(fill="x", anchor="w")
        tk.Label(
            aba,
            text=f'Lida da pasta "{NOME_PASTA_DOCUMENTOS}" (do lado deste programa), arquivo\n'
                 f'"{NOME_PLANILHA_PROCURADORES}". Selecionando uma agência, acrescenta a\n'
                 "assinatura da CREDORA representada pelos procuradores dela (em azul).",
            fg="#555555", justify="left", anchor="w",
        ).pack(fill="x", anchor="w", pady=(4, 0))

    # ------------------------------------------------------------------
    def log(self, mensagem):
        self.log_text.configure(state="normal")
        self.log_text.insert("end", mensagem + "\n")
        self.log_text.see("end")
        self.log_text.configure(state="disabled")
        self.update_idletasks()

    def selecionar_arquivo(self):
        caminho = filedialog.askopenfilename(
            title="Selecione o arquivo .PRN",
            filetypes=[("Arquivos PRN", "*.PRN;*.prn"), ("Todos os arquivos", "*.*")],
        )
        if caminho:
            self.caminho_prn.set(caminho)

    def processar(self):
        caminho = self.caminho_prn.get().strip()
        if not caminho or not os.path.isfile(caminho):
            messagebox.showerror("Erro", "Selecione um arquivo .PRN valido primeiro.")
            return

        self.botao_processar.config(state="disabled")
        self.botao_gerar.config(state="disabled")
        self.log_text.configure(state="normal")
        self.log_text.delete("1.0", "end")
        self.log_text.configure(state="disabled")

        try:
            self.dados_preparados = preparar_arquivo(caminho, log=self.log)
            deteccao = self.dados_preparados["deteccao_imoveis"]

            self.texto_revisao.delete("1.0", "end")
            if deteccao:
                avisos = []
                if deteccao["tem_hipoteca"]:
                    avisos.append("⚠ Hipoteca encontrada na seção de imóveis.")
                if deteccao["avaliacoes_duplicadas"]:
                    avisos.append("⚠ Os dois valores de avaliação do imóvel são IDÊNTICOS -- confira se está certo.")
                if not avisos:
                    avisos.append("Nenhum problema encontrado -- revise/ajuste os valores abaixo se quiser.")
                self.aviso_label.config(text="\n".join(avisos))
                self.texto_revisao.insert("1.0", montar_texto_revisao(deteccao))
            else:
                self.aviso_label.config(text="Documento não tem seção \"2 - IMÓVEIS\" -- nada para revisar aqui.")

            # aba Assinaturas
            self.texto_assinaturas.delete("1.0", "end")
            self.texto_assinaturas.insert("1.0", montar_texto_revisao_assinaturas(self.dados_preparados["assinaturas_qualificadas"]))

            # aba Opcoes: clausula de superveniencia
            tem_clausula = self.dados_preparados["tem_clausula_superveniencia"]
            self.var_superveniencia.set(tem_clausula)
            self.label_status_superveniencia.config(
                text="(encontrada na cédula)" if tem_clausula else "(não encontrada -- marque para acrescentar)"
            )

            # aba Opcoes: cartorio detectado
            cartorio = self.dados_preparados["cartorio_detectado"]
            if cartorio:
                self.combo_cartorio.config(values=[cartorio])
                self.var_cartorio.set(cartorio)
            else:
                self.combo_cartorio.config(values=[])
                self.var_cartorio.set("")

            # aba Opcoes: procuradores (le a planilha "documentos/Procuradores - Cartórios.xlsx")
            self.agencias_procuradores = ler_planilha_procuradores(log=self.log)
            self.combo_procurador.config(values=[a["rotulo"] for a in self.agencias_procuradores])
            self.var_procurador.set("")

            self.log("\nProcessado. Confira as abas acima e clique em \"2. Gerar Word\".")
            self.botao_gerar.config(state="normal")
        except Exception as e:
            self.log("\nERRO: " + str(e))
            self.log(traceback.format_exc())
            messagebox.showerror("Erro ao processar", str(e))
        finally:
            self.botao_processar.config(state="normal")

    def gerar(self):
        if not self.dados_preparados:
            messagebox.showerror("Erro", "Clique em \"1. Processar\" primeiro.")
            return

        self.botao_gerar.config(state="disabled")
        try:
            texto_revisao = None
            if self.dados_preparados["deteccao_imoveis"]:
                texto_revisao = self.texto_revisao.get("1.0", "end-1c")

            texto_assinaturas_editado = self.texto_assinaturas.get("1.0", "end-1c")

            rotulo_agencia = self.var_procurador.get().strip()
            agencia_selecionada = next(
                (a for a in self.agencias_procuradores if a["rotulo"] == rotulo_agencia), None
            )

            pasta_saida = os.path.dirname(self.caminho_prn.get().strip())
            caminho_saida = finalizar_geracao(
                self.dados_preparados,
                pasta_saida,
                texto_revisao_imoveis=texto_revisao,
                checkbox_superveniencia=self.var_superveniencia.get(),
                texto_revisao_assinaturas=texto_assinaturas_editado,
                agencia_procurador=agencia_selecionada,
                log=self.log,
            )
            self.log("\nConcluído com sucesso!")
            resposta = messagebox.askyesno(
                "Concluído", f"Arquivo gerado:\n{caminho_saida}\n\nDeseja abrir a pasta agora?"
            )
            if resposta:
                os.startfile(pasta_saida)  # Windows
        except Exception as e:
            self.log("\nERRO: " + str(e))
            self.log(traceback.format_exc())
            messagebox.showerror("Erro ao gerar", str(e))
        finally:
            self.botao_gerar.config(state="normal")


if __name__ == "__main__":
    if len(sys.argv) > 1:
        # modo linha de comando: python qualificador_app.py caminho.PRN
        # (usa a secao de imoveis tal como veio, sem tela de revisao)
        caminho = sys.argv[1]
        pasta_saida = os.path.dirname(os.path.abspath(caminho))
        dados = preparar_arquivo(caminho)
        finalizar_geracao(dados, pasta_saida)
    else:
        app = App()
        app.mainloop()
